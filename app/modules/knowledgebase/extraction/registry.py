"""
Registry and dispatcher for structured knowledge base file extraction.
"""

from __future__ import annotations

import logging
import re
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document

from app.modules.knowledgebase.extraction.docx_extractor import extract_docx_structured
from app.modules.knowledgebase.extraction.image_extractor import extract_image_structured
from app.modules.knowledgebase.extraction.models import StructuredDocument
from app.modules.knowledgebase.extraction.pdf_extractor import extract_pdf_structured
from app.modules.knowledgebase.extraction.vision import reset_vision_session

logger = logging.getLogger(__name__)

STRUCTURED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".ppt",
    ".pptx",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
}


def normalize_extracted_text(text: str) -> str:
    """Normalize text for storage and future vectorization."""
    cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _extract_legacy_doc_text(file_path: Path) -> str:
    """Extract readable text from a legacy DOC file."""
    try:
        result = subprocess.run(
            ["antiword", str(file_path)],
            capture_output=True,
            text=True,
            timeout=60,
            check=True,
        )
        return normalize_extracted_text(result.stdout)
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
        logger.warning(
            "DOC extraction fallback used for %s; install antiword for better results.",
            file_path,
        )
        raw_bytes = file_path.read_bytes()
        decoded = raw_bytes.decode("utf-8", errors="ignore")
        ascii_text = re.sub(r"[^\x20-\x7E\n]", " ", decoded)
        return normalize_extracted_text(ascii_text)


def _extract_plain_text(file_path: Path) -> StructuredDocument:
    document = StructuredDocument()
    document.add("paragraph", file_path.read_text(encoding="utf-8", errors="ignore"))
    return document


def _extract_csv_structured(file_path: Path) -> StructuredDocument:
    document = StructuredDocument()
    dataframe = pd.read_csv(file_path)
    from app.modules.knowledgebase.extraction.table_utils import dataframe_to_markdown

    document.add("table", dataframe_to_markdown(dataframe))
    return document


def _extract_docx_fallback(file_path: Path) -> StructuredDocument:
    document = StructuredDocument()
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            document.add("paragraph", paragraph.text.strip())
    return document


def extract_structured_document(file_path: Path, file_type: str) -> StructuredDocument:
    """Extract structured content for a supported file type."""
    reset_vision_session()

    normalized_type = file_type.strip().lower()
    if normalized_type and not normalized_type.startswith("."):
        normalized_type = f".{normalized_type}"

    if normalized_type == ".pdf":
        return extract_pdf_structured(file_path)
    if normalized_type == ".docx":
        try:
            return extract_docx_structured(file_path)
        except Exception:
            logger.exception("Structured DOCX extraction failed; using fallback file=%s", file_path)
            return _extract_docx_fallback(file_path)
    if normalized_type in {".ppt", ".pptx"}:
        try:
            from app.modules.knowledgebase.extraction.ppt_extractor import (
                extract_pptx_structured,
            )

            return extract_pptx_structured(file_path)
        except Exception:
            logger.exception("Structured PPT extraction failed file=%s", file_path)
            document = StructuredDocument()
            document.add(
                "paragraph",
                "Presentation uploaded, but structured slide extraction failed for this file.",
            )
            return document
    if normalized_type in {".png", ".jpg", ".jpeg", ".webp"}:
        return extract_image_structured(file_path, normalized_type)
    if normalized_type == ".doc":
        document = StructuredDocument()
        document.add("paragraph", _extract_legacy_doc_text(file_path))
        return document
    if normalized_type in {".txt", ".md"}:
        return _extract_plain_text(file_path)
    if normalized_type == ".csv":
        return _extract_csv_structured(file_path)

    raise ValueError(f"Unsupported file type: {file_type}")


def extract_structured_file_text(file_path: Path, file_type: str) -> str:
    """
    Extract structured content and merge it into searchable plain text.

    Falls back gracefully for unsupported structured types.
    """
    normalized_type = file_type.strip().lower()
    if normalized_type and not normalized_type.startswith("."):
        normalized_type = f".{normalized_type}"

    try:
        structured = extract_structured_document(file_path, normalized_type)
        merged = structured.merge_to_text()
        if merged:
            return normalize_extracted_text(merged)
    except Exception:
        logger.exception(
            "Structured extraction failed file=%s type=%s",
            file_path,
            normalized_type,
        )

    if normalized_type == ".docx":
        return normalize_extracted_text(_extract_docx_fallback(file_path).merge_to_text())
    raise ValueError(f"Unsupported file type: {file_type}")
